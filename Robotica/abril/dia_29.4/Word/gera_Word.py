from docx import Document

#1- Criar um novo documento
documento = Document()


#2- Adicionar um Título principal
documento.add_heading('Relatório de Atividades', level=1)


#3- Adicionar parágrafo com texto
documento.add_paragraph('Este documento foi gerado automaticamente!')
documento.add_paragraph('Novo texto...')


#4- Adicionar um subtitulo
documento.add_heading('Atividades', level=2)


#5- Adicionar a lista de Atividades
atividades = [
    'Reunião com equipe de projetos',
    'Desenvolvimento de novos módulos',
    'Testes de funcionalidade',
    'Treinamento para novos colaboradores'
]

for item in atividades:
    documento.add_paragraph(item, style='List Bullet')


#6- Novo subtitulo de considerações finais
documento.add_heading('Considerações finais', level=2)


#7- paragrafo final
documento.add_paragraph('Todas as metas foram atingidas.')


#8- salvar arquivo
documento.save('Relatorio_atividades.docx')
print('Arquivo criado!')


