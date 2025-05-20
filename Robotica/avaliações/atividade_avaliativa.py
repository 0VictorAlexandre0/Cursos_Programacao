import requests
from deep_translator import GoogleTranslator
import sqlite3
from docx import Document


def coletar_paises():
    pais_pt = []
    for paises in range(3):
        pais = input(f"Digite o {paises+1}º país: ")
        pais_pt.append(pais)
    print(f"Países digitados: {pais_pt}")
    return pais_pt


def traduzir_paises(lista_pt):
    tradutor = GoogleTranslator(source='pt', target='en')
    paises_traduzido = []
    for pais in lista_pt:
        pais_traduzido = tradutor.translate(pais)
        paises_traduzido.append(pais_traduzido)
    return paises_traduzido


def info_pais(pais_en):

    dados_coletados = []

    for pais in pais_en:
        url = f'https://restcountries.com/v3.1/name/{pais}'
        resposta = requests.get(url)

        if resposta.status_code == 200:
            dados = resposta.json()
            info = dados[0]

            print('====================')

            nome_comum = info['name'].get('common', 'Desconhecido')
            print(f'Nome comum: {nome_comum}')

            nome_oficial = info['name'].get('official', 'Desconhecido')
            print(f'Nome oficial: {nome_oficial}')

            capital = info.get('capital', ['Desconhecida'])[0]
            print(f'Capital: {capital}')

            continente = info.get('continents', ['Desconhecido'])[0]
            print(f'Continente: {continente}')

            regiao = info.get('region', 'Desconhecida')
            print(f'Região: {regiao}')
            
            sub_regiao = info.get('subregion', 'Desconhecida')
            print(f'Sub-região: {sub_regiao}')

            populacao = info.get('population', 0)
            print(f'População: {populacao}')

            area = info.get('area', 0.0)
            print(f'Área total: {area} km²')
        
            moedas = info.get('currencies')
            if moedas:
                for dados_moeda in moedas.values():
                    moeda_nome = dados_moeda.get('name', 'Desconhecida')
                    moeda_simbolo = dados_moeda.get('symbol', 'N/A')
                    print(f'Moeda: {moeda_nome} ({moeda_simbolo})')
                    break

            idiomas = info.get('languages')
            if idiomas:
                for nome_idioma in idiomas.values():
                    print(f'Idioma principal: {nome_idioma}')
                    break
            else:
                print('Idioma principal: Informação não disponível')

            fusos = info.get('timezones', ['Sem fuso'])
            for fuso in fusos:
                print(f'Fusos horários: {fuso}')

            bandeiras = info.get('flags')
            if bandeiras:
                url_bandeira = bandeiras.get('png', 'Sem URL')
                print(f'URL da bandeira: {url_bandeira}')
            else:
                print('URL da bandeira: Informação não disponível')


            print('====================')

            conexao_bd(
                nome_comum, nome_oficial, capital, continente, regiao,
                sub_regiao, populacao, area, moeda_nome, moeda_simbolo,
                nome_idioma, url_bandeira
            )

            dados_coletados.append({
                'nome_comum': nome_comum,
                'nome_oficial': nome_oficial,
                'capital': capital,
                'continente': continente,
                'regiao': regiao,
                'sub_regiao': sub_regiao,
                'populacao': populacao,
                'area': area,
                'moeda_nome': moeda_nome,
                'moeda_simbolo': moeda_simbolo,
                'idioma_principal': nome_idioma,
                'fusos': fusos,
                'url_bandeira': url_bandeira
            })
            
        else:
            print("País não encontrado ou erro na requisição.")

    return dados_coletados

def conexao_bd(nome_comum, nome_oficial, capital, continente, regiao, sub_regiao, populacao, area, moeda_nome, moeda_simbolo, idioma_principal, url_bandeira):
    conexao = sqlite3.connect('paises.db')
    cursor = conexao.cursor()

    cursor.execute('''
                    CREATE TABLE IF NOT EXISTS paises(
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        nome_comum TEXT,bra
                        nome_oficial TEXT,
                        capital TEXT,
                        continente TEXT,
                        regiao TEXT,
                        sub_regiao TEXT,
                        populacao INTEGER,
                        area REAL,
                        moeda_nome TEXT,
                        moeda_simbolo TEXT,
                        idioma_principal TEXT,
                        url_bandeira TEXT
                    )
                   ''')
    
    cursor.execute('''
        INSERT INTO paises (nome_comum, nome_oficial, capital, continente, regiao, sub_regiao, populacao, area, moeda_nome, moeda_simbolo, idioma_principal, url_bandeira)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (nome_comum, nome_oficial, capital, continente, regiao, sub_regiao, populacao, area, moeda_nome, moeda_simbolo, idioma_principal, url_bandeira))
    
    conexao.commit()

    print('\nDados inseridos no Banco:')
    cursor.execute('SELECT * FROM paises')
    for linha in cursor.fetchall():
        print(linha)
    
    conexao.close()


def gerar_documento(dados_coletados):
    documento = Document()
    documento.add_heading('Relatório países', level=1)

    for pais in dados_coletados:
        documento.add_heading(pais['nome_comum'], level=2)
        documento.add_paragraph(f"Nome oficial: {pais['nome_oficial']}")
        documento.add_paragraph(f"Capital: {pais['capital']}")
        documento.add_paragraph(f"Continente: {pais['continente']}")
        documento.add_paragraph(f"Região: {pais['regiao']}")
        documento.add_paragraph(f"Sub-região: {pais['sub_regiao']}")
        documento.add_paragraph(f"População: {pais['populacao']}")
        documento.add_paragraph(f"Área: {pais['area']}km²")
        documento.add_paragraph(f"Moeda: {pais['moeda_nome']} ({pais['moeda_simbolo']})")
        documento.add_paragraph(f"Idioma principal: {pais['idioma_principal']}")
        documento.add_paragraph(f"Fusos horários: {pais['fusos']}")
        documento.add_paragraph(f"URL da bandeira: {pais['url_bandeira']}")
        documento.add_paragraph('---')

    documento.save('Relatorio_Paises.docx')
    print('Documento Word gerado.')


def main():
    pais_pt = coletar_paises()
    pais_en = traduzir_paises(pais_pt)
    dados = info_pais(pais_en)
    gerar_documento(dados)

if __name__ == "__main__":
    main()